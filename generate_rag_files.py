import os
from fpdf import FPDF
from docx import Document

# Ensure target directory exists
out_dir = "rag_docs_to_upload"
os.makedirs(out_dir, exist_ok=True)

# ---------------------------------------------------------
# Test Case 1: RAG is Unique (PDF)
# LC Agent has NO text file about IT Security. 
# Only RAG can answer questions about this.
# ---------------------------------------------------------
pdf1 = FPDF()
pdf1.add_page()
pdf1.set_font("Arial", size=12)
pdf1.cell(200, 10, txt="GlobalCorp IT Security Policy v2.0", ln=True, align='C')
pdf1.ln(10)
it_text = (
    "1. PASSWORD REQUIREMENTS\n"
    "All employees must use a 16-character password containing at least one uppercase letter, "
    "one number, and one special character. Passwords must be rotated every 90 days. "
    "Biometric login (Windows Hello or TouchID) is required for all hardware.\n\n"
    "2. VPN USAGE\n"
    "When connecting from public Wi-Fi (e.g., airports, cafes), employees MUST connect to the "
    "GlobalCorp Cisco AnyConnect VPN before accessing any internal domains or reading emails.\n\n"
    "3. PHISHING REPORTING\n"
    "If an employee suspects an email is a phishing attempt, they must click the 'PhishAlarm' "
    "button in Outlook. Forwarding the email to IT Support is no longer the approved method."
)
pdf1.multi_cell(0, 10, txt=it_text)
pdf1.output(os.path.join(out_dir, "IT_Security_Policy.pdf"))

# ---------------------------------------------------------
# Test Case 2: RAG has MORE Context (DOCX)
# LC Agent has a brief leave_policy.txt.
# RAG gets this detailed docx with specific numbers and edge cases.
# ---------------------------------------------------------
doc = Document()
doc.add_heading('GlobalCorp Leave Policy - Comprehensive Handbook', 0)
doc.add_paragraph(
    "1. STANDARD PTO ACCRUAL\n"
    "Full-time salaried employees accrue 20 days of Paid Time Off (PTO) per calendar year. "
    "Employees with 5+ years of tenure accrue 25 days of PTO annually.\n\n"
    "2. PARENTAL LEAVE [EXTENDED DATA]\n"
    "GlobalCorp provides 16 weeks of fully paid parental leave for all new parents (birthing, non-birthing, "
    "and adoptive). In addition, employees returning from parental leave are granted a 'Phase-Back' "
    "period of 4 weeks where they may work 50% capacity at 100% pay.\n\n"
    "3. SABBATICAL LEAVE [EXTENDED DATA]\n"
    "After 7 years of continuous full-time employment, employees are eligible for a 6-week paid sabbatical "
    "to pursue personal projects, volunteering, or rest. Sabbaticals must be approved 6 months in advance "
    "by the VP of the department."
)
doc.save(os.path.join(out_dir, "Leave_Policy_Extended.docx"))

# ---------------------------------------------------------
# Test Case 3: RAG has LESS Context (PDF)
# LC Agent has the highly detailed remote_work_policy.txt.
# RAG only gets this highly abbreviated summary.
# ---------------------------------------------------------
pdf3 = FPDF()
pdf3.add_page()
pdf3.set_font("Arial", size=12)
pdf3.cell(200, 10, txt="Remote Work - Brief Summary", ln=True, align='C')
pdf3.ln(10)
remote_text = (
    "Remote work is generally supported at GlobalCorp for eligible roles. "
    "Employees must discuss remote work arrangements with their direct managers. "
    "While working remotely, employees are expected to maintain professional standards "
    "and ensure a quiet working environment during meetings. "
    "Please refer to the full text policy for details on equipment stipends and core hours."
)
pdf3.multi_cell(0, 10, txt=remote_text)
pdf3.output(os.path.join(out_dir, "Remote_Work_Summary.pdf"))

print(f"Successfully generated Test Case files in '{out_dir}/' directory.")
